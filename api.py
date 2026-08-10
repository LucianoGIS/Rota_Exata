from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import xml.etree.ElementTree as ET
from shapely.geometry import Polygon
import json
import os
import io
from fastapi.responses import Response

# Import the existing solver functions
from mixed_cpp_solver import build_graph_from_polygon, to_working_multidigraph, solve_route, route_to_geojson, export_route_table
import networkx as nx

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class CalculateRequest(BaseModel):
    coordinates: list
    single_pass_twoway: bool = True
    ignore_u_turns: bool = True

@app.post("/upload-kml")
async def upload_kml(file: UploadFile = File(...)):
    contents = await file.read()
    try:
        # Parse KML
        root = ET.fromstring(contents)
        ns = {'kml': 'http://www.opengis.net/kml/2.2'}
        # Find first Polygon coordinates
        coords_node = root.find('.//kml:Polygon//kml:coordinates', ns)
        if coords_node is None:
            # Fallback namespace or lack thereof
            coords_node = root.find('.//Polygon//coordinates')
            
        if coords_node is None:
            raise HTTPException(status_code=400, detail="Nenhum polígono encontrado no KML.")
            
        coords_str = coords_node.text.strip()
        points = []
        for pair in coords_str.split():
            parts = pair.split(',')
            if len(parts) >= 2:
                points.append([float(parts[0]), float(parts[1])]) # lon, lat
                
        if len(points) < 3:
            raise HTTPException(status_code=400, detail="Polígono inválido.")
            
        return {"polygon": points}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


def calculate_bearing_turn(u_prev, u, v, G_nodes):
    if not u_prev:
        return "Início"
        
    u_prev_data = G_nodes.get(u_prev, {})
    u_data = G_nodes.get(u, {})
    v_data = G_nodes.get(v, {})
    
    x_prev, y_prev = u_prev_data.get('x'), u_prev_data.get('y')
    x_u, y_u = u_data.get('x'), u_data.get('y')
    x_v, y_v = v_data.get('x'), v_data.get('y')
    
    if None in (x_prev, y_prev, x_u, y_u, x_v, y_v):
        return "Siga em frente"
        
    dx1, dy1 = x_u - x_prev, y_u - y_prev
    dx2, dy2 = x_v - x_u, y_v - y_u
    
    norm1 = math.hypot(dx1, dy1)
    norm2 = math.hypot(dx2, dy2)
    if norm1 == 0 or norm2 == 0:
        return "Siga em frente"
        
    cross = dx1 * dy2 - dy1 * dx2
    dot = dx1 * dx2 + dy1 * dy2
    
    angle = math.degrees(math.atan2(cross, dot))
    
    if abs(angle) < 25.0:
        return "Siga em frente"
    elif 25.0 <= angle < 120.0:
        return "Vire à esquerda"
    elif -120.0 < angle <= -25.0:
        return "Vire à direita"
    else:
        return "Faça o retorno"


@app.post("/calculate")
async def calculate_route(req: CalculateRequest):
    try:
        polygon_coords = req.coordinates
        polygon = Polygon(polygon_coords)
        
        G_osm = build_graph_from_polygon(polygon)
        G = to_working_multidigraph(G_osm, single_pass_twoway=req.single_pass_twoway, ignore_u_turns=req.ignore_u_turns)
        
        largest_scc = max(nx.strongly_connected_components(G), key=len)
        G = G.subgraph(largest_scc).copy()
        
        start_node = list(G.nodes)[0]
        route = solve_route(G, start_node, start_node)
        
        # Build table data and continuous single LineString coordinates
        table_data = []
        full_route_coords = []
        
        for i, (u, v) in enumerate(zip(route[:-1], route[1:])):
            u_prev = route[i - 1] if i > 0 else None
            direcao = calculate_bearing_turn(u_prev, u, v, G.nodes)

            edge_data = G.get_edge_data(u, v)
            is_reversed = False
            if not edge_data:
                edge_data = G.get_edge_data(v, u)
                is_reversed = True
                
            u_x, u_y = G.nodes[u]['x'], G.nodes[u]['y']
            v_x, v_y = G.nodes[v]['x'], G.nodes[v]['y']
            
            step_coords = []
            if edge_data:
                best_key = min(edge_data, key=lambda k: edge_data[k].get("length", 1.0))
                data = edge_data[best_key]
                
                table_data.append({
                    "passo": i + 1,
                    "direcao": direcao,
                    "rua": data.get("name", "Desconhecida"),
                    "distancia_m": round(data.get("length", 0.0), 2)
                })
                
                if "geometry" in data:
                    coords_list = list(data["geometry"].coords)
                    if is_reversed:
                        coords_list = coords_list[::-1]
                    else:
                        d_start = (coords_list[0][0] - u_x)**2 + (coords_list[0][1] - u_y)**2
                        d_end = (coords_list[-1][0] - u_x)**2 + (coords_list[-1][1] - u_y)**2
                        if d_end < d_start:
                            coords_list = coords_list[::-1]
                    step_coords = coords_list
                else:
                    step_coords = [(u_x, u_y), (v_x, v_y)]
            else:
                # Failsafe: Roteamento físico via G_osm (garante 100% que segue vias reais)
                try:
                    path_nodes = nx.shortest_path(G_osm.to_undirected(), u, v, weight="length")
                    path_dist = 0.0
                    for sub_u, sub_v in zip(path_nodes[:-1], path_nodes[1:]):
                        sub_edge_data = G_osm.get_edge_data(sub_u, sub_v) or G_osm.get_edge_data(sub_v, sub_u)
                        if sub_edge_data:
                            best_sub_k = min(sub_edge_data, key=lambda k: sub_edge_data[k].get("length", 1.0))
                            sub_d = sub_edge_data[best_sub_k]
                            path_dist += sub_d.get("length", 0.0)
                            
                            sub_ux, sub_uy = G_osm.nodes[sub_u]['x'], G_osm.nodes[sub_u]['y']
                            if "geometry" in sub_d:
                                sub_coords = list(sub_d["geometry"].coords)
                                d_s = (sub_coords[0][0] - sub_ux)**2 + (sub_coords[0][1] - sub_uy)**2
                                d_e = (sub_coords[-1][0] - sub_ux)**2 + (sub_coords[-1][1] - sub_uy)**2
                                if d_e < d_s:
                                    sub_coords = sub_coords[::-1]
                                step_coords.extend(sub_coords)
                            else:
                                step_coords.extend([(sub_ux, sub_uy), (G_osm.nodes[sub_v]['x'], G_osm.nodes[sub_v]['y'])])
                    
                    table_data.append({
                        "passo": i + 1,
                        "direcao": direcao,
                        "rua": "Conexão de Vias",
                        "distancia_m": round(path_dist, 2)
                    })
                except Exception:
                    table_data.append({
                        "passo": i + 1,
                        "direcao": direcao,
                        "rua": "Desconhecida",
                        "distancia_m": 0.0
                    })
                    step_coords = [(u_x, u_y), (v_x, v_y)]

            # Concatenar step_coords em full_route_coords sem duplicar vértices adjacentes
            if not full_route_coords:
                full_route_coords.extend(step_coords)
            else:
                if step_coords:
                    if full_route_coords[-1] == step_coords[0]:
                        full_route_coords.extend(step_coords[1:])
                    else:
                        full_route_coords.extend(step_coords)
            
        geojson = {
            "type": "FeatureCollection",
            "features": [{
                "type": "Feature",
                "geometry": {"type": "LineString", "coordinates": full_route_coords},
                "properties": {
                    "n_nodes": len(route),
                    "n_edges": len(route) - 1,
                }
            }]
        }
        
        return {"geojson": geojson, "table": table_data}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class ExportRequest(BaseModel):
    table: list
    geojson: dict = None

@app.post("/export/pdf")
async def export_pdf(req: ExportRequest):
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(190, 10, txt="Relatório de Rota Otimizada", ln=1, align='C')
    pdf.ln(5)
    
    pdf.set_font("Arial", 'B', 9)
    pdf.cell(15, 8, "#", 1, 0, 'C')
    pdf.cell(40, 8, "Direção", 1, 0, 'L')
    pdf.cell(105, 8, "Rua / Trecho", 1, 0, 'L')
    pdf.cell(30, 8, "Distância (m)", 1, 1, 'R')
    
    pdf.set_font("Arial", '', 9)
    for row in req.table:
        pdf.cell(15, 7, str(row['passo']), 1, 0, 'C')
        pdf.cell(40, 7, str(row.get('direcao', 'Siga em frente')), 1, 0, 'L')
        pdf.cell(105, 7, str(row['rua'])[:50], 1, 0, 'L')
        pdf.cell(30, 7, str(row['distancia_m']), 1, 1, 'R')
        
    pdf_bytes = bytes(pdf.output())
    return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=rota.pdf"})

@app.post("/export/docx")
async def export_docx(req: ExportRequest):
    from docx import Document
    doc = Document()
    doc.add_heading('Relatório de Rota Otimizada', 0)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Passo'
    hdr_cells[1].text = 'Direção / Manobra'
    hdr_cells[2].text = 'Rua / Trecho'
    hdr_cells[3].text = 'Distância (m)'
    
    for row in req.table:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['passo'])
        row_cells[1].text = str(row.get('direcao', 'Siga em frente'))
        row_cells[2].text = str(row['rua'])
        row_cells[3].text = str(row['distancia_m'])
        
    f = io.BytesIO()
    doc.save(f)
    return Response(content=f.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=rota.docx"})

@app.post("/export/html")
async def export_html(req: ExportRequest):
    html = "<html><head><meta charset='utf-8'><title>Relatório de Rota Otimizada</title>"
    html += "<style>"
    html += "body{font-family: Arial, sans-serif; margin: 20px; background:#f9f9f9; color:#333;}"
    html += "h2{color:#2b5a9e; border-bottom:2px solid #98c43d; padding-bottom:10px;}"
    html += "table{border-collapse:collapse; width:100%; background:#fff; box-shadow:0 1px 3px rgba(0,0,0,0.1); border-radius:6px; overflow:hidden;}"
    html += "th,td{border:1px solid #e2e8f0; padding:10px 12px; text-align:left; font-size:14px;}"
    html += "th{background:#f1f5f9; color:#475569; font-weight:600;}"
    html += "tr:nth-child(even){background:#f8fafc;}"
    html += ".badge{display:inline-block; padding:4px 8px; border-radius:4px; font-weight:600; font-size:12px;}"
    html += ".badge-direta{background:#e0f2fe; color:#0369a1;}"
    html += ".badge-esquerda{background:#f0fdf4; color:#15803d;}"
    html += ".badge-frente{background:#f1f5f9; color:#475569;}"
    html += ".badge-retorno{background:#fef3c7; color:#b45309;}"
    html += ".badge-inicio{background:#f3e8ff; color:#6b21a8;}"
    html += "</style></head><body>"
    html += "<h2>Relatório de Rota Otimizada por Ruas</h2>"
    html += "<table><tr><th style='width:60px;'>Passo</th><th style='width:160px;'>Direção / Manobra</th><th>Rua / Trecho</th><th style='width:120px;text-align:right;'>Distância (m)</th></tr>"
    for row in req.table:
        dir_text = str(row.get('direcao', 'Siga em frente'))
        badge_cls = "badge-frente"
        if "direita" in dir_text.lower(): badge_cls = "badge-direta"
        elif "esquerda" in dir_text.lower(): badge_cls = "badge-esquerda"
        elif "retorno" in dir_text.lower(): badge_cls = "badge-retorno"
        elif "início" in dir_text.lower() or "inicio" in dir_text.lower(): badge_cls = "badge-inicio"
        
        html += f"<tr><td><b>#{row['passo']}</b></td><td><span class='badge {badge_cls}'>{dir_text}</span></td><td>{row['rua']}</td><td style='text-align:right;'>{row['distancia_m']} m</td></tr>"
    html += "</table></body></html>"
    return Response(content=html, media_type="text/html", headers={"Content-Disposition": "attachment; filename=rota.html"})

@app.post("/export/kml")
async def export_kml(req: ExportRequest):
    import simplekml
    kml = simplekml.Kml()
    geom_type = req.geojson['features'][0]['geometry']['type']
    coords = req.geojson['features'][0]['geometry']['coordinates']
    
    if geom_type == 'MultiLineString':
        # Create a folder to group the lines
        fol = kml.newfolder(name="Rota Otimizada")
        for i, line_coords in enumerate(coords):
            fol.newlinestring(name=f"Trecho {i+1}", description="Caminho do veículo", coords=line_coords)
    else:
        kml.newlinestring(name="Rota Otimizada", description="Caminho do veículo", coords=coords)
        
    kml_str = kml.kml()
    return Response(content=kml_str, media_type="application/vnd.google-earth.kml+xml", headers={"Content-Disposition": "attachment; filename=rota.kml"})
